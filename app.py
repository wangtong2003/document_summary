import pymysql
pymysql.install_as_MySQLdb()

from flask import Flask, request, jsonify, Response, render_template, stream_with_context, send_file, make_response, session
from flask_session import Session  # 添加 Flask-Session 导入
import asyncio
from ollama import Client
import os
import PyPDF2
from docx import Document
import markdown
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import mysql.connector
import uuid
from datetime import datetime, timedelta
import jwt
from functools import wraps
from werkzeug.utils import secure_filename
from flask_jwt_extended import (
    JWTManager, jwt_required, get_jwt_identity, create_access_token,
    exceptions as jwt_exceptions
)
import json
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import shutil
import requests
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import jieba
from flask_migrate import Migrate
from io import BytesIO
from urllib.parse import quote
import traceback

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
DOCUMENTS_FOLDER = 'documents'  # 新增永久文档存储目录
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'md', 'epub'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOCUMENTS_FOLDER'] = DOCUMENTS_FOLDER  # 新增配置
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json; charset=utf8mb4'

# MySQL配置
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:20030221@localhost/document_summary?charset=utf8mb4'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 10,
    'pool_recycle': 3600,
    'pool_pre_ping': True,
    'pool_timeout': 30,
    'max_overflow': 5
}

# Session 配置
app.config['SECRET_KEY'] = 'your-secret-key-here'  # 设置 session 密钥
app.config['SESSION_TYPE'] = 'filesystem'  # 使用文件系统存储 session
app.config['SESSION_FILE_DIR'] = 'flask_session'  # session 文件存储目录
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)  # session 有效期

# 初始化 Flask-Session
Session(app)

db = SQLAlchemy(app)
migrate = Migrate(app, db)  # 初始化 Flask-Migrate

# 创建必要的目录
for folder in [UPLOAD_FOLDER, DOCUMENTS_FOLDER, 'flask_session']:
    if not os.path.exists(folder):
        os.makedirs(folder)

# 数据模型定义
class User(db.Model):
    """用户模型"""
    __tablename__ = 'users'
    
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)  # 添加自增属性
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    role = db.Column(db.String(20), nullable=False, default='user')
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.now)
    
    def __repr__(self):
        return f'<User {self.username}>'

class DocumentSummary(db.Model):
    __tablename__ = 'document_summaries'
    
    id = db.Column(db.Integer, primary_key=True)
    file_name = db.Column(db.String(255), nullable=False)
    file_hash = db.Column(db.String(32), nullable=False)
    summary_text = db.Column(db.Text(length=4294967295), nullable=False)
    original_text = db.Column(db.Text(length=4294967295), nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    summary_length = db.Column(db.String(20))
    target_language = db.Column(db.String(20))
    file_content = db.Column(db.LargeBinary(length=(2**32)-1))
    file_size = db.Column(db.BigInteger)
    mime_type = db.Column(db.String(100))
    original_filename = db.Column(db.String(255))
    display_filename = db.Column(db.String(255))
    keywords = db.Column(db.String(255))  # 新增关键字字段

    __table_args__ = {
        'mysql_engine': 'InnoDB',
        'mysql_charset': 'utf8mb4',
        'mysql_collate': 'utf8mb4_unicode_ci'
    }

class FileMapping(db.Model):
    __tablename__ = 'file_mappings'
    
    id = db.Column(db.Integer, primary_key=True)
    summary_id = db.Column(db.Integer, db.ForeignKey('document_summaries.id', ondelete='CASCADE'), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    system_filename = db.Column(db.String(255), nullable=False)
    display_filename = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    summary = db.relationship('DocumentSummary', backref=db.backref('file_mappings', lazy=True))

def save_summary_to_db(file_info, summary_text, params, file_content=None):
    """保存或更新文档摘要到MySQL"""
    try:
        print("\n=== 开始保存摘要到数据库 ===")
        print(f"文件信息: {file_info}")
        print(f"摘要长度: {len(summary_text) if summary_text else 0}")
        print(f"参数: {params}")
        print(f"关键字: {file_info.get('keywords', '无')}")
        
        # 计算文件内容的MD5哈希值
        import hashlib
        file_hash = hashlib.md5(file_info["original_text"].encode('utf-8')).hexdigest()
        
        # 获取原始文件名和显示文件名
        original_filename = file_info.get("original_filename")
        display_filename = original_filename  # 使用原始文件名作为显示文件名
        
        print(f"原始文件名: {original_filename}")
        print(f"显示文件名: {display_filename}")
        
        # 检查是否已存在相同文件的摘要
        existing_summary = DocumentSummary.query.filter_by(
            file_hash=file_hash
        ).first()
        
        try:
            if existing_summary:
                print(f"更新现有摘要 ID: {existing_summary.id}")
                # 更新现有摘要
                existing_summary.summary_text = summary_text
                existing_summary.summary_length = params.get("summary_length")
                existing_summary.target_language = params.get("target_language")
                existing_summary.original_filename = original_filename
                existing_summary.display_filename = display_filename
                existing_summary.keywords = file_info.get('keywords')  # 更新关键字
                if file_content:
                    existing_summary.file_content = file_content
                    existing_summary.file_size = len(file_content)
                    existing_summary.mime_type = file_info.get('mime_type')
                existing_summary.updated_at = datetime.now()
                
                # 更新文件名映射
                file_mapping = FileMapping.query.filter_by(summary_id=existing_summary.id).first()
                if file_mapping:
                    file_mapping.original_filename = original_filename
                    file_mapping.system_filename = file_info["filename"]
                    file_mapping.display_filename = display_filename
                else:
                    new_mapping = FileMapping(
                        summary_id=existing_summary.id,
                        original_filename=original_filename,
                        system_filename=file_info["filename"],
                        display_filename=display_filename
                    )
                    db.session.add(new_mapping)
                
                db.session.commit()
                print("摘要更新成功")
            else:
                print("创建新摘要记录")
                # 创建新摘要记录
                new_summary = DocumentSummary(
                    file_name=file_info["filename"],
                    file_hash=file_hash,
                    summary_text=summary_text,
                    original_text=file_info["original_text"],
                    summary_length=params.get("summary_length"),
                    target_language=params.get("target_language"),
                    file_content=file_content,
                    file_size=len(file_content) if file_content else None,
                    mime_type=file_info.get('mime_type'),
                    original_filename=original_filename,
                    display_filename=display_filename,
                    keywords=file_info.get('keywords')  # 保存关键字
                )
                db.session.add(new_summary)
                db.session.flush()  # 获取新插入记录的ID
                
                # 创建文件名映射
                new_mapping = FileMapping(
                    summary_id=new_summary.id,
                    original_filename=original_filename,
                    system_filename=file_info["filename"],
                    display_filename=display_filename
                )
                db.session.add(new_mapping)
                db.session.commit()
                print("摘要保存成功")
            
            return True
            
        except Exception as e:
            print(f"数据库操作出错: {str(e)}")
            print(f"完整错误信息: {e.__class__.__name__}: {str(e)}")
            traceback.print_exc()
            db.session.rollback()
            raise
            
    except Exception as e:
        print(f"保存摘要到数据库时出错: {str(e)}")
        print(f"完整错误信息: {e.__class__.__name__}: {str(e)}")
        traceback.print_exc()
        db.session.rollback()
        raise

# 文档处理函数
def read_pdf(file_path):
    """读取PDF文件内容"""
    try:
        print(f"开始读取PDF文件: {file_path}")
        text = ""
        with open(file_path, 'rb') as file:
            # 创建PDF文件阅读器对象
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 获取PDF页数
            num_pages = len(pdf_reader.pages)
            print(f"PDF文件共有 {num_pages} 页")
            
            # 遍历每一页并提取文本
            for page_num in range(num_pages):
                try:
                    # 获取页面对象
                    page = pdf_reader.pages[page_num]
                    # 提取文本
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    print(f"成功读取第 {page_num + 1} 页，提取到 {len(page_text)} 个字符")
                except Exception as e:
                    print(f"读取第 {page_num + 1} 页时出错: {str(e)}")
                    continue
            
            if not text.strip():
                raise Exception("未能从PDF中提取任何文本")
                
            print(f"PDF文件读取完成，总共提取到 {len(text)} 个字符")
            return text
    except Exception as e:
        print(f"读取PDF文件出错: {str(e)}")
        raise Exception(f"读取PDF文件失败: {str(e)}")

def read_docx(file_path):
    doc = Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def read_txt(file_path):
    """读取文本文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # 如果 UTF-8 失败，尝试其他编码
        encodings = ['gbk', 'gb2312', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise Exception("无法识别文件编码")

def read_md(file_path):
    """读取 Markdown 文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return markdown.markdown(text)
    except UnicodeDecodeError:
        # 果 UTF-8 失败，尝试其他编码
        encodings = ['gbk', 'gb2312', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                return markdown.markdown(text)
            except UnicodeDecodeError:
                continue
        raise Exception("无法识别文编码")

def read_epub(file_path):
    book = epub.read_epub(file_path)
    text = ''
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_body_content(), 'html.parser')
            text += soup.get_text() + '\n'
    return text

def read_document(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    if file_extension == '.pdf':
        return read_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:  # 同时支持 .docx 和 .doc
        return read_docx(file_path)
    elif file_extension == '.txt':
        return read_txt(file_path)
    elif file_extension == '.md':
        return read_md(file_path)
    elif file_extension == '.epub':
        return read_epub(file_path)
    else:
        raise ValueError("不支持的文件格式")

def generate_summary(text):
    """生成文档摘要"""
    try:
        print("开始生成摘要...")
        print(f"输入文本长度: {len(text)} 字符")
        
        client = Client(host='http://localhost:11434')
        
        # 构建提示词
        system_prompt = """You are a powerful AI assistant capable of reading text content and summarizing it. 
        Please follow these basic requirements:

        [BASIC REQUIREMENTS]
        - Summary Length: 约500字
        - Target Language: chinese
        - Keywords: Generate exactly 4 keywords that best represent the main topics of the text

        [OUTPUT FORMAT]
        Please format your response as follows:

        [KEYWORDS]
        keyword1|keyword2|keyword3|keyword4

        [SUMMARY]
        Your summary text here...

        [INSTRUCTIONS]
        1. First output exactly 4 keywords separated by | character
        2. Then output the summary after [SUMMARY] marker
        3. Keywords should be concise and representative
        4. Write everything in chinese
        """
        
        content = f"""
        ARTICLE_TO_SUMMARY:
        <ARTICLE>
        {text}
        </ARTICLE>
        """
        
        print("调用程 Ollama API...")
        response = client.generate(
            model='qwen2.5:3b',
            prompt=system_prompt + content,
            stream=False
        )
        
        if not response or 'response' not in response:
            raise Exception("API 返回的数据格式不正确")
            
        summary = response['response']
        print(f"成功生成摘要，长度: {len(summary)} 字符")
        return summary
        
    except Exception as e:
        print(f"生成摘要时发生错误: {str(e)}")
        raise Exception(f"生成摘要失败: {str(e)}")

def ollama_text(input_text, params=None, file_info=None, file_content=None):
    try:
        client = Client(host='http://localhost:11434')
        
        # 根据参数调整提示词
        summary_length_map = {
            'very_short': '约100字',
            'short': '约200字',
            'medium': '约500字',
            'long': '约1000字',
            'very_long': '约2000字'
        }
        
        # 获取参数
        summary_length = params.get('summary_length', 'medium')
        target_length = summary_length_map.get(summary_length, '约500字')
        target_language = params.get('target_language', 'chinese')
        
        # 构建基础提示模板
        system_prompt = f"""你是一个专业的文档摘要专家。请仔细阅读以下内容，并生成结构化摘要和关键字。

        [输出格式要求]
        请严格按照以下格式输出：
        1. 首先输出[KEYWORDS]标记
        2. 然后在新行输出4个关键字，用|分隔
        3. 再输出[SUMMARY]标记
        4. 最后输出摘要内容

        [基本要求]
        - 摘要长度：{target_length}
        - 目标语言：{target_language}
        - 关键字数量：必须exactly输出4个关键字
        - 关键字要简洁且具有代表性
        - 所有内容使用中文输出

        [示例格式]
        [KEYWORDS]
        关键词1|关键词2|关键词3|关键词4
        [SUMMARY]
        这里是摘要正文内容...
        """
        
        content = f"""
        需要总结的文档内容：
        <文档>
        {input_text}
        </文档>
        """
        
        print("调用远程 Ollama API...")
        # 使用非流式响应以确保完整性
        response = client.generate(
            model='qwen2.5:3b',
            prompt=system_prompt + content,
            stream=False
        )
        
        if not response or 'response' not in response:
            raise Exception("API 返回的数据格式不正确")
            
        full_response = response['response']
        print("收到完整响应：", full_response[:100], "...")  # 打印响应的前100个字符
        
        # 解析关键字和摘要
        keywords = None
        summary_text = None
        
        # 分离关键字和摘要
        try:
            # 使用更严格的解析逻辑
            if '[KEYWORDS]' in full_response and '[SUMMARY]' in full_response:
                # 提取关键字部分
                keywords_section = full_response.split('[KEYWORDS]')[1].split('[SUMMARY]')[0].strip()
                # 提取摘要部分
                summary_text = full_response.split('[SUMMARY]')[1].strip()
                
                # 确保关键字格式正确
                if '|' in keywords_section:
                    keywords = keywords_section
                    print(f"成功提取关键字: {keywords}")
                else:
                    print("关键字格式不正确，尝试修复...")
                    # 尝试清理和规范化关键字
                    cleaned_keywords = keywords_section.replace('\n', '').replace('，', '|').replace(',', '|')
                    if cleaned_keywords:
                        keywords = cleaned_keywords
                        print(f"修复后的关键字: {keywords}")
            else:
                print("响应中未找到关键字或摘要标记")
        except Exception as e:
            print(f"解析关键字时出错: {str(e)}")
            keywords = None
        
        if not summary_text:
            raise Exception("无法提取摘要内容")
            
        print(f"提取的关键字: {keywords}")
        print(f"提取的摘要(前100字): {summary_text[:100]}...")
        
        # 保存到数据库
        if file_info:
            try:
                if keywords:
                    file_info['keywords'] = keywords
                save_result = save_summary_to_db(file_info, summary_text, params, file_content)
                if not save_result:
                    print("警告：摘要保存失败")
            except Exception as e:
                print(f"保存摘要失败: {str(e)}")
                traceback.print_exc()
        
        # 返回流式响应
        def generate():
            # 首先输出关键字部分
            if keywords:
                yield f"data: [KEYWORDS]\n"
                yield f"data: {keywords}\n\n"
            
            # 然后输出摘要部分
            yield f"data: [SUMMARY]\n"
            # 将摘要分段输出
            for i in range(0, len(summary_text), 100):
                chunk = summary_text[i:i+100]
                yield f"data: {chunk}\n\n"
        
        return Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'Connection': 'keep-alive',
                'X-Accel-Buffering': 'no'
            }
        )
        
    except Exception as e:
        print(f"Ollama API错误: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def init_admin():
    """初始化管理员账户"""
    try:
        admin = User.query.filter_by(username='admin').first()
        if not admin:
            admin = User(
                username='admin',
                email='admin@example.com',
                password=generate_password_hash('admin', method='pbkdf2:sha256'),  # 使用正确的哈希方法
                role='admin',
                created_at=datetime.now()
            )
            db.session.add(admin)
            db.session.commit()
            print("管理员账户初始化成功")
        else:
            # 更新管理员密码以确保使用正确的哈希方法
            admin.password = generate_password_hash('admin', method='pbkdf2:sha256')
            db.session.commit()
            print("管理员账户已存在，已更新密码哈希")
    except Exception as e:
        print(f"初始化管理员账户失败: {str(e)}")
        db.session.rollback()

@app.route('/login', methods=['POST'])
def login():
    """用户登录"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({'error': '用户名和密码不能为空'}), 400
            
        user = User.query.filter_by(username=username).first()
        
        if not user:
            return jsonify({'error': '用户名或密码错误'}), 401
            
        if check_password_hash(user.password, password):
            # 登录成功
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            
            return jsonify({
                'message': '登录成功',
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'role': user.role
                }
            })
        else:
            return jsonify({'error': '用户名密码错误'}), 401
            
    except Exception as e:
        print(f"登录失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/logout')
def logout():
    """用户登出"""
    session.clear()
    return jsonify({'message': '登出成功'})

@app.route('/api/check_auth')
def check_auth():
    """检查用户认证状态"""
    if 'user_id' in session:
        return jsonify({
            'authenticated': True,
            'user': {
                'id': session['user_id'],
                'username': session['username'],
                'role': session['role']
            }
        })
    return jsonify({'authenticated': False}), 401

@app.route('/')
def index():
    """主页"""
    return render_template('dashboard.html')

@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/register')
def register_page():
    return render_template('register.html')

@app.route('/summary_library')
def summary_library():
    """染摘要库页面"""
    return render_template('summaries.html')

@app.route('/summaries')
def get_summaries():
    """获取所有摘要列表"""
    try:
        print("\n=== 开始获取摘要列表 ===")
        summaries = DocumentSummary.query.order_by(DocumentSummary.created_at.desc()).all()
        print(f"查询到 {len(summaries)} 条摘要记录")
        
        results = []
        for summary in summaries:
            # 使用原始文件名作为显示名称
            display_name = summary.original_filename or summary.display_filename or summary.file_name
            
            # 获取文件扩展名
            file_type = os.path.splitext(display_name)[1].lower().lstrip('.') if display_name else 'unknown'
            
            # 打印详细的调试信息
            print(f"\n处理摘要 ID: {summary.id}")
            print(f"文件名: {display_name}")
            print(f"关键字: {summary.keywords}")
            print(f"摘要文本: {summary.summary_text[:100]}...")
            
            # 构建结果数据
            result_data = {
                'id': summary.id,
                'file_name': display_name,
                'file_type': file_type,
                'summary_text': summary.summary_text,
                'created_at': summary.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'summary_length': summary.summary_length,
                'target_language': summary.target_language,
                'file_size': summary.file_size,
                'mime_type': summary.mime_type,
                'has_file': bool(summary.file_content),
                'keywords': summary.keywords
            }
            
            print(f"构建的结果数据:")
            print(f"- ID: {result_data['id']}")
            print(f"- 文件名: {result_data['file_name']}")
            print(f"- 关键字: {result_data['keywords']}")
            
            results.append(result_data)
            
        print(f"\n返回 {len(results)} 条摘要记录")
        return jsonify(results)
        
    except Exception as e:
        print(f"获取摘要列表错误: {str(e)}")
        traceback.print_exc()  # 打印完整的错误堆栈
        return jsonify({'error': str(e)}), 500

@app.route('/summaries/<int:summary_id>', methods=['GET'])
def get_summary_detail(summary_id):
    """获取单摘要详情"""
    try:
        print(f"\n=== 获取摘要详情 ID: {summary_id} ===")
        summary = DocumentSummary.query.get(summary_id)
        
        if not summary:
            print(f"未找到ID为 {summary_id} 的摘要")
            return jsonify({'error': f'未找到ID为 {summary_id} 的摘要'}), 404
            
        result = {
            'id': summary.id,
            'file_name': summary.file_name,
            'summary_text': summary.summary_text,
            'original_text': summary.original_text,
            'created_at': summary.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'summary_length': summary.summary_length,
            'target_language': summary.target_language
        }
        return jsonify(result)
        
    except Exception as e:
        print(f"获取摘要详情错误: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/summaries/<int:summary_id>', methods=['DELETE'])
def delete_summary(summary_id):
    """删除指定摘要"""
    try:
        print(f"\n=== 开始删除摘要 ID: {summary_id} ===")
        
        # 首先查找摘要记录
        summary = DocumentSummary.query.get(summary_id)
        if not summary:
            error_msg = f"未找到ID为 {summary_id} 的摘要记录"
            print(error_msg)
            return jsonify({'error': error_msg}), 404
            
        print(f"找到摘要记录: {summary.file_name}")
        
        try:
            # 删除关联的文件映射记录
            mappings = FileMapping.query.filter_by(summary_id=summary_id).all()
            if mappings:
                print(f"删除 {len(mappings)} 个关联的文件映射记录")
                for mapping in mappings:
                    db.session.delete(mapping)
            else:
                print("没有找到关联的文件映射记录")
            
            # 删除摘要记录
            print("删除摘要记录")
            db.session.delete(summary)
            
            # 提交事务
            db.session.commit()
            print(f"摘要删除成功")
            
            return jsonify({'message': '删除成功'})
            
        except Exception as e:
            print(f"删除过程中出错: {str(e)}")
            db.session.rollback()
            return jsonify({'error': f'删除失败: {str(e)}'}), 500
            
    except Exception as e:
        print(f"删除摘要时出错: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/favicon.ico')
def favicon():
    return '', 204

def sanitize_filename(filename):
    """安全地处理文件名，保中文字符"""
    # 移除文件名中的特殊字符，但保留中文字符
    name, ext = os.path.splitext(filename)
    # 只保留中文、英文、数字和下划
    name = re.sub(r'[^\w\u4e00-\u9fff-]', '_', name)
    # 确保文件名不空
    if not name:
        name = 'document'
    return name + ext

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    extension = os.path.splitext(filename)[1].lower()
    print(f"检查文件扩展名: {extension}")
    print(f"允许的扩展名: {ALLOWED_EXTENSIONS}")
    # 移除扩展名前的点号再检查
    return extension.lstrip('.') in ALLOWED_EXTENSIONS

@app.route('/process_document', methods=['POST'])
def process_document():
    """处理上传的文档并生成摘要"""
    try:
        print("\n=== 开始处理文档 ===")
        
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
            
        params = {
            'summary_length': request.form.get('summary_length', 'medium'),
            'target_language': request.form.get('target_language', 'chinese'),
            'summary_style': request.form.get('summary_style'),
            'focus_area': request.form.get('focus_area'),
            'expertise_level': request.form.get('expertise_level'),
            'language_style': request.form.get('language_style')
        }
        
        # 保存原始文件名（不进行任何处理）
        original_filename = file.filename
        print(f"原始文件名: {original_filename}")
        
        # 检查文件类型
        if not allowed_file(original_filename):
            print(f"文件扩展名检查失败: {original_filename}")
            return jsonify({'error': '不支持的文件格式'}), 400
        
        # 生成系统文件名（用于内部存储
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        _, ext = os.path.splitext(original_filename)  # 从原始文件名获取扩展名
        system_filename = f"document_{timestamp}{ext}"  # 使用通用前缀
        print(f"系统文件名: {system_filename}")
        
        # 读取文件内容
        file_content = file.read()
        file.seek(0)  # 重置文件指针
        
        # 创建临时文件用于读取内容
        temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], system_filename)
        file.save(temp_file_path)
        
        try:
            # 读取文档内容
            text = read_document(temp_file_path)
            
            # 准备文件信息
            file_info = {
                'filename': system_filename,  # 系统内部使用的文件名
                'original_filename': original_filename,  # 原始文件名（未经过任何处理）
                'original_text': text,
                'mime_type': get_file_mime_type(original_filename)
            }
            
            # 生成摘要并保存
            return ollama_text(text, params, file_info, file_content)
            
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                
    except Exception as e:
        print(f"处理文档时发生错误: {str(e)}")
        return jsonify({'error': f'处理文档失败: {str(e)}'}), 500

class ChineseTokenizer:
    def __call__(self, text):
        return list(jieba.cut(text))

def search_summaries(query, summaries, top_k=5):
    """
    用TF-IDF和余弦相似度进行智能检索
    """
    if not summaries:
        return []
        
    # 准备文档集合
    documents = [summary.summary_text for summary in summaries]
    documents.append(query)  # 将查询添加到文档集合中
    
    # 创建TF-IDF向量化器，使用自定义的中文分词器
    vectorizer = TfidfVectorizer(
        tokenizer=ChineseTokenizer(),
        stop_words='english',  # 移除英文停用词
        max_features=5000  # 限制特征数量
    )
    
    try:
        # 转换文档为TF-IDF矩阵
        tfidf_matrix = vectorizer.fit_transform(documents)
        
        # 计算查询（最后一个文档）与其他文档的余弦相似度
        cosine_similarities = cosine_similarity(
            tfidf_matrix[-1:], tfidf_matrix[:-1]
        ).flatten()
        
        # 获取相似度最高的档索引
        top_indices = np.argsort(cosine_similarities)[::-1][:top_k]
        
        # 构建结果列表
        results = []
        for idx in top_indices:
            if cosine_similarities[idx] > 0:  # 只返回相似度大于0的结果
                summary = summaries[idx]
                results.append({
                    'id': summary.id,
                    'file_name': summary.original_filename or summary.display_filename or summary.file_name,
                    'summary_text': summary.summary_text,
                    'created_at': summary.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'similarity_score': float(cosine_similarities[idx]),
                    'summary_length': summary.summary_length,
                    'target_language': summary.target_language,
                    'original_filename': summary.original_filename,
                    'display_filename': summary.display_filename
                })
        
        return results
    except Exception as e:
        print(f"搜索过程中出错: {str(e)}")
        return []

@app.route('/api/search', methods=['POST'])
def search_api():
    """智能检索API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        
        if not query:
            return jsonify({'error': '搜索查询不能为空'}), 400
            
        # 获取所有摘要
        summaries = DocumentSummary.query.all()
        
        # 执行智能检索
        results = search_summaries(query, summaries)
        
        # 确保使用原始文件名
        for result in results:
            result['file_name'] = result.get('original_filename') or result.get('display_filename') or result.get('file_name')
        
        return jsonify({
            'results': results,
            'total': len(results)
        })
        
    except Exception as e:
        print(f"搜索API错误: {str(e)}")
        return jsonify({'error': str(e)}), 500

def get_file_mime_type(filename):
    """获取件的MIME类型"""
    mime_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.epub': 'application/epub+zip'
    }
    ext = os.path.splitext(filename)[1].lower()
    return mime_types.get(ext, 'application/octet-stream')

def save_uploaded_file(file, filename):
    """保存上传的文件久存储目录"""
    # 生成唯一的文件名
    unique_filename = f"{str(uuid.uuid4())}{os.path.splitext(filename)[1]}"
    file_path = os.path.join(app.config['DOCUMENTS_FOLDER'], unique_filename)
    file.save(file_path)
    return file_path, unique_filename

def convert_docx_to_html(file_path):
    """将Word文档转换为HTML"""
    try:
        doc = Document(file_path)
        html_content = ['<!DOCTYPE html><html><head><meta charset="UTF-8"><style>body{font-family:Arial,sans-serif;line-height:1.6;padding:20px;}</style></head><body>']
        
        for paragraph in doc.paragraphs:
            if paragraph.text:
                style = paragraph.style.name
                if style.startswith('Heading'):
                    level = style[-1]
                    html_content.append(f'<h{level}>{paragraph.text}</h{level}>')
                else:
                    html_content.append(f'<p>{paragraph.text}</p>')
        
        for table in doc.tables:
            html_content.append('<table border="1" style="border-collapse:collapse;margin:10px 0;">')
            for row in table.rows:
                html_content.append('<tr>')
                for cell in row.cells:
                    html_content.append(f'<td style="padding:8px;">{cell.text}</td>')
                html_content.append('</tr>')
            html_content.append('</table>')
            
        html_content.append('</body></html>')
        return '\n'.join(html_content)
    except Exception as e:
        print(f"转换Word文档时出错: {str(e)}")
        raise

def convert_epub_to_html(file_path):
    """将EPUB转换为HTML"""
    try:
        book = epub.read_epub(file_path)
        html_content = ['<!DOCTYPE html><html><head><meta charset="UTF-8"><style>body{font-family:Arial,sans-serif;line-height:1.6;padding:20px;}</style></head><body>']
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                # 提取并保留原有的HTML构
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'img']):
                    html_content.append(str(element))
        
        html_content.append('</body></html>')
        return '\n'.join(html_content)
    except Exception as e:
        print(f"转换EPUB出错: {str(e)}")
        raise

@app.route('/preview/<int:summary_id>')
def preview_file(summary_id):
    """预览始文档"""
    try:
        summary = DocumentSummary.query.get_or_404(summary_id)
        
        if not summary.file_content:
            return jsonify({'error': '找不到文件内容'}), 404
            
        # 使用原始文件名作为显示名称
        display_name = summary.original_filename or summary.display_filename or summary.file_name
        print(f"预览文件: {display_name}")
            
        # 创建内存文件对象
        file_stream = BytesIO(summary.file_content)
        
        # 根据文件类型处理预览
        if summary.mime_type == 'application/pdf':
            response = send_file(
                file_stream,
                mimetype='application/pdf',
                as_attachment=False,
                download_name=display_name
            )
            # 添加响应头，确保文件名正确显示
            response.headers['Content-Disposition'] = f'inline; filename*=UTF-8\'\'{quote(display_name)}'
            response.headers['Content-Type'] = 'application/pdf'
            return response
            
        elif summary.mime_type in ['text/plain', 'text/markdown']:
            # 对于文本文件，直接读取内容
            text_content = file_stream.read().decode('utf-8')
            if summary.mime_type == 'text/markdown':
                html_content = markdown.markdown(text_content)
                response = make_response(f'''
                <!DOCTYPE html>
                <html lang="zh-CN">
                <head>
                    <meta charset="UTF-8">
                    <title>{display_name}</title>
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            line-height: 1.6;
                            padding: 20px;
                            max-width: 800px;
                            margin: 0 auto;
                        }}
                    </style>
                </head>
                <body>
                    <h1>{display_name}</h1>
                    {html_content}
                </body>
                </html>
                ''')
                response.headers['Content-Type'] = 'text/html; charset=utf-8'
                return response
            else:
                response = make_response(f'''
                <!DOCTYPE html>
                <html lang="zh-CN">
                <head>
                    <meta charset="UTF-8">
                    <title>{display_name}</title>
                    <style>
                        body {{
                            font-family: monospace;
                            white-space: pre-wrap;
                            padding: 20px;
                            max-width: 800px;
                            margin: 0 auto;
                        }}
                    </style>
                </head>
                <body>
                    <h1>{display_name}</h1>
                    <pre>{text_content}</pre>
                </body>
                </html>
                ''')
                response.headers['Content-Type'] = 'text/html; charset=utf-8'
                return response
            
        elif summary.mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            try:
                # 创建临时文件
                temp_file = BytesIO(summary.file_content)
                doc = Document(temp_file)
                
                html_content = [f'''<!DOCTYPE html>
                <html lang="zh-CN">
                <head>
                    <meta charset="UTF-8">
                    <title>{display_name}</title>
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            line-height: 1.6;
                            padding: 20px;
                            max-width: 800px;
                            margin: 0 auto;
                        }}
                    </style>
                </head>
                <body>
                    <h1>{display_name}</h1>
                ''']
                
                for paragraph in doc.paragraphs:
                    if paragraph.text:
                        style = paragraph.style.name
                        if style.startswith('Heading'):
                            level = style[-1]
                            html_content.append(f'<h{level}>{paragraph.text}</h{level}>')
                        else:
                            html_content.append(f'<p>{paragraph.text}</p>')
                
                for table in doc.tables:
                    html_content.append('<table border="1" style="border-collapse:collapse;margin:10px 0;width:100%;">')
                    for row in table.rows:
                        html_content.append('<tr>')
                        for cell in row.cells:
                            html_content.append(f'<td style="padding:8px;">{cell.text}</td>')
                        html_content.append('</tr>')
                    html_content.append('</table>')
                    
                html_content.append('</body></html>')
                response = make_response('\n'.join(html_content))
                response.headers['Content-Type'] = 'text/html; charset=utf-8'
                return response
            except Exception as e:
                print(f"Word文档转换错误: {str(e)}")
                return jsonify({'error': '无法转换Word文档'}), 500
            
        elif summary.mime_type == 'application/epub+zip':
            try:
                # 创建临时文件
                temp_file = BytesIO(summary.file_content)
                book = epub.read_epub(temp_file)
                
                html_content = [f'''<!DOCTYPE html>
                <html lang="zh-CN">
                <head>
                    <meta charset="UTF-8">
                    <title>{display_name}</title>
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            line-height: 1.6;
                            padding: 20px;
                            max-width: 800px;
                            margin: 0 auto;
                        }}
                    </style>
                </head>
                <body>
                    <h1>{display_name}</h1>
                ''']
                
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'img']):
                            html_content.append(str(element))
                
                html_content.append('</body></html>')
                response = make_response('\n'.join(html_content))
                response.headers['Content-Type'] = 'text/html; charset=utf-8'
                return response
            except Exception as e:
                print(f"EPUB转换错误: {str(e)}")
                return jsonify({'error': '无法转换EPUB文件'}), 500
            
        else:
            return jsonify({
                'message': '此类型文件不支持在线预览',
                'file_type': summary.mime_type
            }), 400
            
    except Exception as e:
        print(f"预览文件时出错: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<int:summary_id>')
def download_file(summary_id):
    """下载原始文档"""
    try:
        summary = DocumentSummary.query.get_or_404(summary_id)
        
        if not summary.file_content:
            return jsonify({'error': '找不到文件内容'}), 404
            
        # 使用原始文件名或显示文件名
        display_name = summary.original_filename or summary.display_filename or summary.file_name
            
        # 创建临时文件用于下载
        temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], summary.file_name)
        with open(temp_file_path, 'wb') as f:
            f.write(summary.file_content)
        
        try:
            return send_file(
                temp_file_path,
                mimetype=summary.mime_type,
                as_attachment=True,
                download_name=display_name
            )
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
        
    except Exception as e:
        print(f"下载文件时出错: {str(e)}")
        return jsonify({'error': str(e)}), 500

def migrate_existing_data():
    """迁移现有数据，保所有必要的字段都存在"""
    try:
        with app.app_context():
            summaries = DocumentSummary.query.all()
            for summary in summaries:
                if not hasattr(summary, 'file_content'):
                    summary.file_content = None
                if not hasattr(summary, 'file_size'):
                    summary.file_size = None
                if not hasattr(summary, 'mime_type'):
                    summary.mime_type = get_file_mime_type(summary.file_name)
                if not hasattr(summary, 'updated_at'):
                    summary.updated_at = summary.created_at
            db.session.commit()
            print("数据迁移完成")
    except Exception as e:
        print(f"数据迁移失败: {str(e)}")
        db.session.rollback()

@app.route('/admin/users')
def admin_users():
    """用户管理页面"""
    return render_template('admin/users.html')

@app.route('/api/users', methods=['GET'])
def get_users():
    """获取用户列表"""
    try:
        users = User.query.all()
        return jsonify([{
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'role': user.role,
            'created_at': user.created_at.strftime('%Y-%m-%d %H:%M:%S')
        } for user in users])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/users/<user_id>', methods=['DELETE'])
def delete_user(user_id):
    """删除用户"""
    try:
        user = User.query.get_or_404(user_id)
        if user.username == 'admin':
            return jsonify({'error': '不能删除管理员账户'}), 403
        db.session.delete(user)
        db.session.commit()
        return jsonify({'message': '户删除成功'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/users/<user_id>', methods=['PUT'])
def update_user(user_id):
    """更新用户信息"""
    try:
        user = User.query.get_or_404(user_id)
        data = request.get_json()
        
        if 'username' in data and data['username'] != user.username:
            if User.query.filter_by(username=data['username']).first():
                return jsonify({'error': '用户名已存在'}), 400
            user.username = data['username']
            
        if 'email' in data and data['email'] != user.email:
            if User.query.filter_by(email=data['email']).first():
                return jsonify({'error': '邮箱已存在'}), 400
            user.email = data['email']
            
        if 'role' in data and data['role'] != user.role:
            if user.username == 'admin' and data['role'] != 'admin':
                return jsonify({'error': '不能修改管理员角色'}), 403
            user.role = data['role']
            
        if 'password' in data and data['password']:
            user.password = generate_password_hash(data['password'], method='pbkdf2:sha256')
            
        db.session.commit()
        return jsonify({
            'message': '户信息更新成功',
            'user': {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'role': user.role,
                'created_at': user.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/users/<user_id>', methods=['GET'])
def get_user(user_id):
    """获取单个用户信息"""
    try:
        user = User.query.get_or_404(user_id)
        return jsonify({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'role': user.role,
            'created_at': user.created_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        # 只在表不存在时创建表
        db.create_all()
        # 初始化管理员账户
        init_admin()
        print("数据库初始化完成")
    app.run(debug=True, port=5000)
